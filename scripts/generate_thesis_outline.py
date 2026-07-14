"""Generate thesis outline .docx following CTU-QD1799 structure."""
import sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pathlib import Path

OUTPUT = Path(__file__).parent.parent / "docs" / "THESIS_OUTLINE_CTU1799.docx"
doc = Document()

# -- Page setup --
for section in doc.sections:
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(2.0)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)

style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(13)
style.paragraph_format.line_spacing = 1.2

def add_heading_ctu(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = 'Times New Roman'
        run.font.size = Pt(14) if level <= 1 else Pt(13)
    return h

def add_status(text, done=True):
    p = doc.add_paragraph()
    tag = "✅ ĐÃ CÓ" if done else "❌ CẦN BỔ SUNG"
    run = p.add_run(f"[{tag}] ")
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0, 128, 0) if done else RGBColor(200, 0, 0)
    run.bold = True
    run2 = p.add_run(text)
    run2.font.size = Pt(12)
    run2.font.name = 'Times New Roman'
    return p

def add_note(text):
    p = doc.add_paragraph()
    run = p.add_run(f"    → Ghi chú: {text}")
    run.font.size = Pt(11)
    run.font.italic = True
    run.font.color.rgb = RGBColor(100, 100, 100)

# ============================================================
# TITLE
# ============================================================
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title.add_run("OUTLINE LUẬN VĂN THẠC SĨ\nTheo Chuẩn CTU-QD1799")
r.bold = True
r.font.size = Pt(18)
r.font.name = 'Times New Roman'

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = sub.add_run("Đề tài: Dự Báo Nồng Độ Bụi Mịn PM2.5 Sử Dụng Học Máy Và Học Sâu\nTrên Dữ Liệu Cảm Biến IoT Tại Đồng Bằng Sông Cửu Long")
r2.font.size = Pt(14)
r2.font.name = 'Times New Roman'

doc.add_paragraph("─" * 60)

# ============================================================
# PHẦN MỞ ĐẦU (Trang bìa, tóm tắt, etc.)
# ============================================================
add_heading_ctu("PHẦN MỞ ĐẦU", 1)

add_status("1. Trang bìa chính (Phụ lục 1a)", False)
add_note("Cần soạn: Tên trường, Khoa, Tên tác giả, Tên đề tài, Ngành, Năm")

add_status("2. Trang bìa phụ (Phụ lục 2a)", False)
add_note("Thêm: Mã số HV, Tên CBHD")

add_status("3. Trang xác nhận Hội đồng (Phụ lục 3a)", False)

add_status("4. Lời cảm ơn", False)

add_status("5. Tóm tắt tiếng Việt (200-350 từ)", False)
add_note("Nội dung: (i) chủ đề + mục tiêu, (ii) phương pháp, (iii) kết quả, (iv) kết luận. Từ khóa ≤6 từ")

add_status("6. Tóm tắt tiếng Anh (Abstract)", False)

add_status("7. Lời cam đoan", False)

add_status("8. Mục lục", False)
add_note("Tự động sinh từ Word. Tối đa đến tiểu mục cấp 2 (3 chữ số, VD: 1.2.3)")

add_status("9. Danh sách bảng", False)
add_note("Dashboard có ~30+ bảng/chart, cần đánh số theo chương (VD: Bảng 4.1)")

add_status("10. Danh sách hình", False)

add_status("11. Danh mục từ viết tắt", True)
add_note("Đã có trong Dashboard: PM2.5, MASE, SHAP, IoT, GRU, LSTM, TFT, EDA, STL, ADF, KPSS, S-ESD...")

# ============================================================
# CHƯƠNG 1: GIỚI THIỆU
# ============================================================
add_heading_ctu("CHƯƠNG 1: GIỚI THIỆU", 1)

add_status("1.1 Đặt vấn đề (Bối cảnh ô nhiễm không khí tại ĐBSCL)", True)
add_note("Dashboard Tổng Quan đã có insight cards, bối cảnh PM2.5 ở Sa Đéc")

add_status("1.2 Mục tiêu nghiên cứu", True)
add_note("(1) Xây dựng pipeline ML/DL dự báo PM2.5 multi-horizon; (2) So sánh 30+ mô hình; (3) Giải thích mô hình bằng SHAP")

add_status("1.3 Câu hỏi nghiên cứu", True)
add_note("Q1: Độ phân giải nào tối ưu? Q2: ML hay DL tốt hơn? Q3: Feature nào quan trọng nhất?")

add_status("1.4 Giả thuyết nghiên cứu", False)
add_note("CẦN VIẾT: VD: H1: Ensemble sẽ tốt hơn single model; H2: 30m resolution là tối ưu")

add_status("1.5 Giới hạn và phạm vi nghiên cứu", True)
add_note("Đã xác định: 1 trạm IoT (Sa Đéc), 2022-2025, 5 biến, 3 horizons (1h/6h/24h)")

add_status("1.6 Ý nghĩa khoa học và thực tiễn", False)
add_note("CẦN VIẾT: Đóng góp cho lĩnh vực giám sát CLKK tại ĐBSCL")

# ============================================================
# CHƯƠNG 2: TỔNG QUAN TÀI LIỆU
# ============================================================
add_heading_ctu("CHƯƠNG 2: TỔNG QUAN TÀI LIỆU", 1)

add_status("2.1 Tổng quan về ô nhiễm PM2.5", False)
add_note("CẦN VIẾT: Định nghĩa, nguồn phát thải, tác động sức khỏe, tiêu chuẩn WHO/QCVN")

add_status("2.2 Các phương pháp dự báo chất lượng không khí", True)
add_note("Dashboard citations.py đã có 14 bài báo SOTA verified (2022-2025)")

add_status("2.2.1 Phương pháp thống kê truyền thống (ARIMA, SARIMA)", True)
add_note("Đã thí nghiệm và loại bỏ do dữ liệu non-stationary + gaps lớn")

add_status("2.2.2 Machine Learning (LightGBM, XGBoost, RF, SVR)", True)
add_note("30+ models đã train, kết quả có trong snapshot v9_multi_resolution")

add_status("2.2.3 Deep Learning (GRU, LSTM, TFT)", True)
add_note("Fair vs Expert Pipeline ablation study đã hoàn thành")

add_status("2.2.4 Ensemble Methods (Stacking, Weighted)", True)

add_status("2.3 Các nghiên cứu liên quan tại Việt Nam và quốc tế", True)
add_note("Bảng so sánh literature đã có trong dashboard_content.json: 6 quốc tế + 2 VN")

add_status("2.4 Khoảng trống nghiên cứu (Research Gap)", False)
add_note("CẦN VIẾT: Thiếu nghiên cứu multi-resolution + multi-horizon ở ĐBSCL; Thiếu explainability")

# ============================================================
# CHƯƠNG 3: PHƯƠNG PHÁP NGHIÊN CỨU
# ============================================================
add_heading_ctu("CHƯƠNG 3: PHƯƠNG PHÁP NGHIÊN CỨU", 1)

add_status("3.1 Mô tả dữ liệu và thiết bị IoT", True)
add_note("209K records, 5 biến (PM2.5, Nhiệt độ, Độ ẩm, Điểm sương, CO2), tọa độ 10.29°N, 105.77°E")

add_status("3.1.1 Hạn chế dữ liệu (Data Sparsity)", True)
add_note("ĐÃ PHÂN TÍCH: 89 ngày mù, Tháng 2 & 9 gần trắng. EDA Tab 2 có Blind Spot Matrix")

add_status("3.1.2 Dữ liệu ngoại lai bổ trợ (Open-Meteo)", True)
add_note("File: dataset/external/open_meteo_missing_periods.csv — 20,520 rows hourly data")

add_status("3.2 Quy trình tiền xử lý dữ liệu (Data Pipeline)", True)
add_note("Raw → Clean (S-ESD) → Resample (15m/30m/1h) → Impute (Spline+KNN) → Features (119)")

add_status("3.2.1 Làm sạch dữ liệu (S-ESD Outlier Detection)", True)
add_status("3.2.2 Tiered Imputation Strategy", True)
add_note("Spline (<6h) → KNN Multivariate (6-24h) → Drop (>24h). Có trích dẫn Moritz 2015")

add_status("3.2.3 Feature Engineering (119 features)", True)
add_note("Lag, Rolling, EWM, Fourier, DateTime, Interaction features. Anti-leakage shift(1)")

add_status("3.3 Chiến lược chia dữ liệu (Train/Val/Test Split)", True)
add_note("80/10/10 temporal split. Test-on-Real-Only (is_imputed==0)")

add_status("3.4 Các mô hình dự báo", True)
add_status("3.4.1 Persistence Baseline", True)
add_status("3.4.2 ML Models (LightGBM, XGBoost, RF, SVR, ElasticNet)", True)
add_status("3.4.3 DL Models (GRU, LSTM, TFT)", True)
add_status("3.4.4 Ensemble (Stacking, Weighted Averaging)", True)

add_status("3.5 Metrics đánh giá", True)
add_note("MAE, RMSE, R², MASE (chính, theo Hyndman & Koehler 2006), Brier Score, F1")

add_status("3.6 Phương pháp giải thích mô hình (Explainability)", True)
add_note("SHAP + Permutation Importance. Explainability Hub đã hoàn chỉnh")

add_status("3.7 Kiến trúc hệ thống Dashboard", True)
add_note("Streamlit + FastAPI + PostgreSQL. Docker-ready. VTF theme framework")

# ============================================================
# CHƯƠNG 4: KẾT QUẢ VÀ THẢO LUẬN
# ============================================================
add_heading_ctu("CHƯƠNG 4: KẾT QUẢ VÀ THẢO LUẬN", 1)

add_status("4.1 Kết quả EDA (Khám phá dữ liệu)", True)
add_note("6 tabs EDA đầy đủ: Overview, Gaps, Stationarity, Autocorrelation, Why, Deep Insights")

add_status("4.1.1 Thống kê mô tả đa độ phân giải", True)
add_status("4.1.2 Kiểm định tính dừng (ADF + KPSS dual test)", True)
add_status("4.1.3 Phân tách STL (Trend, Seasonal, Residual)", True)
add_status("4.1.4 Phân tích tự tương quan và Bẫy Persistence", True)
add_status("4.1.5 Granger Causality giữa biến khí tượng và PM2.5", True)

add_status("4.2 Kết quả huấn luyện mô hình", True)
add_note("30+ models × 3 resolutions × 3 horizons. Snapshot v9_multi_resolution")

add_status("4.2.1 So sánh hiệu suất theo Horizon (1h, 6h, 24h)", True)
add_note("Bảng ranking MASE đầy đủ trong Dashboard Overview")

add_status("4.2.2 Thí nghiệm Multi-Resolution Ablation", True)
add_note("15m vs 30m vs 1h. Kết luận: 30m là optimal trade-off")

add_status("4.2.3 Ablation Study: Fair DL vs Expert DL", True)
add_note("Fair Pipeline (tabular features) > Expert Pipeline (raw) cho IoT data")

add_status("4.2.4 Diebold-Mariano Test (Ý nghĩa thống kê)", True)
add_note("DM test xác nhận p < 0.001 cho các best models vs Persistence")

add_status("4.3 Giải thích mô hình (Explainability)", True)
add_status("4.3.1 SHAP Feature Importance", True)
add_note("pm25_lag_1 dominates ở 1h; Fourier features nổi bật ở 24h")
add_status("4.3.2 Permutation Importance", True)

add_status("4.4 Đánh giá Conformal Prediction (Uncertainty)", True)
add_note("Conformal Prediction intervals cho các mô hình best. Brier Score evaluation")

add_status("4.5 Thảo luận", False)
add_note("CẦN VIẾT: So sánh kết quả với literature (Bảng 2.3). Giải thích tại sao Ensemble thắng, tại sao TFT_1h fail, ý nghĩa thực tiễn")

add_status("4.6 Hạn chế của nghiên cứu", True)
add_note("Đã tổng hợp: Data sparsity, đơn trạm, Apple Silicon, chưa online learning")

# ============================================================
# CHƯƠNG 5: KẾT LUẬN VÀ ĐỀ XUẤT
# ============================================================
add_heading_ctu("CHƯƠNG 5: KẾT LUẬN VÀ ĐỀ XUẤT", 1)

add_status("5.1 Kết luận", False)
add_note("CẦN VIẾT: Tóm tắt các kết quả chính liên hệ với mục tiêu nghiên cứu (Chương 1)")

add_status("5.2 Đóng góp của luận văn", False)
add_note("CẦN VIẾT: (1) Pipeline anti-leakage, (2) Multi-resolution methodology, (3) Dashboard công cụ")

add_status("5.3 Đề xuất hướng nghiên cứu tiếp theo", False)
add_note("Gợi ý: Multi-station, Online Learning, Satellite data integration, CNN-BiLSTM, Transfer Learning")

# ============================================================
# PHẦN CUỐI
# ============================================================
add_heading_ctu("PHẦN CUỐI", 1)

add_status("Tài liệu tham khảo (IEEE format)", True)
add_note("14 bài đã verified trong citations.py. CẦN kiểm tra format IEEE đúng chuẩn")

add_status("Phụ lục A: Bảng kết quả chi tiết toàn bộ models", True)
add_note("Có thể export từ Dashboard ranking table")

add_status("Phụ lục B: Mã nguồn chính (Code snippets)", True)

add_status("Phụ lục C: Cấu hình Hyperparameters", True)

# ============================================================
# SUMMARY TABLE
# ============================================================
doc.add_page_break()
add_heading_ctu("BẢNG TỔNG HỢP: ĐÃ CÓ vs CẦN BỔ SUNG", 1)

table = doc.add_table(rows=1, cols=3)
table.style = 'Light Grid Accent 1'
hdr = table.rows[0].cells
hdr[0].text = "STT"
hdr[1].text = "Nội dung"
hdr[2].text = "Trạng thái"

gaps = [
    ("1", "Trang bìa chính + phụ", "❌ Chưa soạn"),
    ("2", "Tóm tắt TV + TA", "❌ Chưa viết"),
    ("3", "Lời cảm ơn + Cam đoan", "❌ Chưa viết"),
    ("4", "Giả thuyết nghiên cứu (1.4)", "❌ Cần bổ sung"),
    ("5", "Ý nghĩa khoa học (1.6)", "❌ Cần viết"),
    ("6", "Tổng quan PM2.5 (2.1)", "❌ Cần viết lý thuyết"),
    ("7", "Research Gap (2.4)", "❌ Cần phân tích"),
    ("8", "Thảo luận kết quả (4.5)", "❌ Cần viết chi tiết"),
    ("9", "Kết luận (5.1)", "❌ Cần tổng hợp"),
    ("10", "Đóng góp (5.2)", "❌ Cần viết"),
    ("11", "Đề xuất (5.3)", "❌ Cần viết"),
    ("12", "Format IEEE references", "⚠️ Cần kiểm tra"),
]

for stt, nd, tt in gaps:
    row = table.add_row().cells
    row[0].text = stt
    row[1].text = nd
    row[2].text = tt

# Save
doc.save(str(OUTPUT))
print(f"✅ Đã xuất file: {OUTPUT}")
print(f"   Tổng: 12 mục CẦN BỔ SUNG, phần còn lại ĐÃ CÓ dữ liệu từ dự án")
